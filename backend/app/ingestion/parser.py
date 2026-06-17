"""Document parser supporting PDF, DOCX, and Markdown formats.

V2: Outputs structured DocumentSection tree in addition to flat text.
Sections preserve heading hierarchy, page ranges, and parent-child relationships.
"""

import re
from pathlib import Path
from dataclasses import dataclass, field


# ────────────────────────────────────────────
# Data models
# ────────────────────────────────────────────

@dataclass
class DocumentSection:
    """A logical section within a document, preserving structural hierarchy."""
    heading: str                          # e.g. "3. Refund Conditions"
    heading_level: int                    # 0 = no heading, 1 = h1, 2 = h2, 3 = h3
    content: str                          # Full text under this section
    page_number: int | None = None
    page_number_end: int | None = None
    subsections: list["DocumentSection"] = field(default_factory=list)

    @property
    def is_leaf(self) -> bool:
        return len(self.subsections) == 0

    def iter_leaves(self):
        """Depth-first iteration over leaf sections."""
        if self.is_leaf:
            yield self
        else:
            for sub in self.subsections:
                yield from sub.iter_leaves()


@dataclass
class ParsedDocument:
    """Result of parsing a document file."""
    text: str                             # Full flat text (backward-compatible)
    sections: list[DocumentSection]       # Top-level sections (new, may be empty)
    page_count: int = 1
    title: str | None = None


# ────────────────────────────────────────────
# Utility: build section path string
# ────────────────────────────────────────────

def build_section_path(section: DocumentSection, ancestors: list[DocumentSection] | None = None) -> str:
    """Build a breadcrumb path like 'Policies > Refunds > Conditions'.

    Args:
        section: The target section.
        ancestors: List of ancestor sections (for internal nodes, passed during tree walk).

    Returns:
        ' > '-separated heading path, or '' if no headings.
    """
    parts: list[str] = []
    if ancestors:
        for a in ancestors:
            if a.heading:
                parts.append(a.heading)
    if section.heading:
        parts.append(section.heading)
    return " > ".join(parts)


# ────────────────────────────────────────────
# PDF Parser
# ────────────────────────────────────────────

# Heuristic patterns for detecting headings in PDF text
_HEADING_PATTERNS: list[re.Pattern] = [
    re.compile(r'^第[一二三四五六七八九十百千\d]+[章节条款部篇].*$'),   # "第三条 定义", "第二章 退款"
    re.compile(r'^\d+[\.、\s]\s*\S.*$'),                              # "3. xxx" or "3、xxx"
    re.compile(r'^\d+\.\d+[\.\s]\s*\S.*$'),                           # "3.1. xxx"
    re.compile(r'^[A-Z][A-Z\s]+\.\s*\S.*$'),                          # "ANNEX A. xxx"
    re.compile(r'^[IVX]+\.\s+\S.*$'),                                  # "IV. xxx"
]


def _looks_like_heading(line: str, font_size: float | None = None, is_bold: bool = False) -> bool:
    """Heuristic heading detection for PDF lines."""
    stripped = line.strip()
    if not stripped or len(stripped) > 120:
        return False

    # Short bold lines are likely headings
    if is_bold and len(stripped) <= 80:
        return True

    # Pattern-based detection
    for pat in _HEADING_PATTERNS:
        if pat.match(stripped):
            return True

    return False


def _estimate_heading_level(line: str, font_size: float | None = None) -> int:
    """Estimate heading level from content pattern and font size."""
    stripped = line.strip()

    # Numbered pattern depth: "3.1.2" → deeper nesting
    dots = stripped.count('.')
    if dots > 0 and re.match(r'^\d+(\.\d+)*[\.\s]', stripped):
        return min(dots + 1, 3)

    # Chinese chapter/article patterns → h1/h2
    if re.match(r'^第[一二三四五六七八九十百千\d]+[章篇部]', stripped):
        return 1
    if re.match(r'^第[一二三四五六七八九十百千\d]+[节条款]', stripped):
        return 2

    # Bold short lines → likely h1 or h2
    if font_size and font_size >= 16:
        return 1
    if font_size and font_size >= 13:
        return 2

    return 3


def parse_pdf(file_path: str) -> ParsedDocument:
    """Extract text and sections from a PDF file using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    title = None
    if reader.metadata and reader.metadata.title:
        title = reader.metadata.title

    all_pages: list[str] = []
    sections: list[DocumentSection] = []
    current_section_lines: list[str] = []
    current_heading = ""
    current_level = 0
    current_page_start: int | None = None

    for page_idx, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if not page_text:
            continue
        all_pages.append(page_text)

        # Try to extract font info for heading detection
        lines = page_text.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # Heuristic heading check
            if _looks_like_heading(stripped):
                # Save previous section
                if current_section_lines:
                    sections.append(DocumentSection(
                        heading=current_heading,
                        heading_level=current_level,
                        content='\n'.join(current_section_lines).strip(),
                        page_number=current_page_start,
                        page_number_end=page_idx + 1,
                    ))
                current_heading = stripped
                current_level = _estimate_heading_level(stripped)
                current_section_lines = []
                current_page_start = page_idx + 1
            else:
                current_section_lines.append(line)

    # Save final section
    if current_section_lines or current_heading:
        sections.append(DocumentSection(
            heading=current_heading,
            heading_level=current_level,
            content='\n'.join(current_section_lines).strip() if current_section_lines else current_heading,
            page_number=current_page_start,
            page_number_end=len(reader.pages),
        ))

    # If no sections detected, create a single root section
    if not sections:
        full_text = "\n\n".join(all_pages)
        sections = [DocumentSection(
            heading="",
            heading_level=0,
            content=full_text,
            page_number=1,
            page_number_end=len(reader.pages),
        )]

    return ParsedDocument(
        text="\n\n".join(all_pages),
        sections=sections,
        page_count=len(reader.pages),
        title=title,
    )


# ────────────────────────────────────────────
# DOCX Parser
# ────────────────────────────────────────────

def _get_heading_level(style_name: str) -> int:
    """Extract heading level from a Word style name like 'Heading 2'."""
    match = re.search(r'(\d+)', style_name)
    if match:
        return int(match.group(1))
    if 'Heading' in style_name or 'heading' in style_name.lower():
        return 1
    return 0


def parse_docx(file_path: str) -> ParsedDocument:
    """Extract text and structured sections from a DOCX file."""
    from docx import Document

    doc = Document(file_path)
    all_paragraphs: list[str] = []
    sections: list[DocumentSection] = []
    current_heading = ""
    current_level = 0
    current_lines: list[str] = []

    title = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Empty line = paragraph break in current section
            if current_lines:
                current_lines.append("")
            continue

        style_name = para.style.name if para.style else ""
        level = _get_heading_level(style_name)

        if level > 0:
            # Save previous section
            if current_lines:
                sections.append(DocumentSection(
                    heading=current_heading,
                    heading_level=current_level,
                    content='\n'.join(current_lines).strip(),
                ))
            elif current_heading and not current_lines:
                # Heading with no body yet — keep as placeholder
                pass

            current_heading = text
            current_level = level
            current_lines = []

            if title is None and level == 1:
                title = text
        else:
            current_lines.append(text)

        all_paragraphs.append(text)

    # Save final section
    if current_lines or current_heading:
        sections.append(DocumentSection(
            heading=current_heading,
            heading_level=current_level,
            content='\n'.join(current_lines).strip() if current_lines else current_heading,
        ))

    if not sections:
        sections = [DocumentSection(
            heading="",
            heading_level=0,
            content="\n\n".join(all_paragraphs),
        )]

    return ParsedDocument(
        text="\n\n".join(all_paragraphs),
        sections=sections,
        page_count=1,
        title=title,
    )


# ────────────────────────────────────────────
# Markdown Parser
# ────────────────────────────────────────────

_MD_HEADING_RE = re.compile(r'^(#{1,6})\s+(.+)$', re.MULTILINE)


def parse_markdown(file_path: str) -> ParsedDocument:
    """Read a Markdown file and extract structured sections from headings."""
    path = Path(file_path)
    text = path.read_text(encoding="utf-8")

    title = None
    lines = text.split('\n')

    # Detect title from first # heading
    for line in lines:
        m = _MD_HEADING_RE.match(line.strip())
        if m and len(m.group(1)) == 1:
            title = m.group(2)
            break

    sections = _parse_md_sections(lines)
    if not sections:
        sections = [DocumentSection(
            heading="",
            heading_level=0,
            content=text,
        )]

    return ParsedDocument(
        text=text,
        sections=sections,
        page_count=1,
        title=title,
    )


def _parse_md_sections(lines: list[str]) -> list[DocumentSection]:
    """Parse Markdown lines into a section tree using heading markers."""
    # Collect (line_index, level, heading_text)
    heading_positions: list[tuple[int, int, str]] = []
    for i, line in enumerate(lines):
        m = _MD_HEADING_RE.match(line.strip())
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading_positions.append((i, level, text))

    if not heading_positions:
        return []

    sections: list[DocumentSection] = []
    # Stack of (level, heading, content_start_line, section_object)
    stack: list[tuple[int, str, int, DocumentSection]] = []

    for idx, (line_idx, level, heading) in enumerate(heading_positions):
        # Determine content end line
        if idx + 1 < len(heading_positions):
            content_end = heading_positions[idx + 1][0]
        else:
            content_end = len(lines)

        # Extract content: lines after heading until next heading
        content_lines = lines[line_idx + 1:content_end]
        content = '\n'.join(content_lines).strip()

        # Pop stack until we find a parent (level < current level)
        while stack and stack[-1][0] >= level:
            stack.pop()

        section = DocumentSection(
            heading=heading,
            heading_level=level,
            content=content,
        )

        if stack:
            # Add as subsection to parent
            stack[-1][3].subsections.append(section)
        else:
            sections.append(section)

        stack.append((level, heading, line_idx, section))

    return sections


# ────────────────────────────────────────────
# Parser registry
# ────────────────────────────────────────────

PARSER_MAP = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "markdown": parse_markdown,
}


def parse_document(file_path: str, doc_type: str) -> ParsedDocument:
    """Parse a document file and return extracted text with section structure.

    Args:
        file_path: Path to the document file.
        doc_type: One of 'pdf', 'docx', 'markdown'.

    Returns:
        ParsedDocument with text, structured sections, and metadata.

    Raises:
        ValueError: If doc_type is not supported.
    """
    parser = PARSER_MAP.get(doc_type)
    if parser is None:
        raise ValueError(f"Unsupported document type: {doc_type}. Supported: {list(PARSER_MAP.keys())}")
    return parser(file_path)
